"""DOCX generator - Convert article markdown to Word document.

Parses markdown-like article text and produces a styled .docx file using
python-docx.  Supports headings, block quotes, bold inline runs, horizontal
rules, interactive question blocks, embedded images, and image URL appendix.
"""
import os
import re
import datetime
import tempfile
import httpx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _set_document_styles(doc):
    """Apply template-matched styles (Arial-based, Chinese-compatible)."""
    # Normal style: Arial 12pt, 1.15 line spacing
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)

    # Heading 1 style: Arial 20pt Bold Dark Blue (for 【】sections)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Arial'
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    h1.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.line_spacing = 1.15

    # Heading 2 style: Arial 14pt Bold (for sub-headings)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Arial'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h2.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.15

def _sanitize_title(title):
    """Strip characters that are illegal in Windows / Unix filenames."""
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '_', title)
    cleaned = cleaned.strip()
    if len(cleaned) > 80:
        cleaned = cleaned[:80]
    return cleaned or "untitled"


def _build_filename(title):
    """Return  YYYYMMDD_<sanitized-title>.docx ."""
    date_part = datetime.datetime.now().strftime("%Y%m%d")
    safe = _sanitize_title(title)
    return date_part + "_" + safe + ".docx"


def _download_image_bytes(url):
    """Attempt to download an image; return bytes or None on any failure."""
    try:
        with httpx.Client(timeout=15, follow_redirects=True) as client:
            resp = client.get(url)
            if resp.status_code == 200 and len(resp.content) > 0:
                return resp.content
    except Exception:
        pass
    return None


def _add_inline_formatted_paragraph(doc, text):
    """Add a paragraph whose **bold** spans become bold runs.

    Supports an arbitrary number of bold segments interleaved with normal
    text.  For example::

        Some **bold** and **more bold** text.

    will produce five runs: normal, bold, normal, bold, normal.
    """
    para = doc.add_paragraph()
    tokens = re.split(r'\*\*(.+?)\*\*', text)
    # tokens follow the pattern: normal, bold_content, normal, bold_content, ...
    for idx, token in enumerate(tokens):
        if not token:
            continue
        run = para.add_run(token)
        if idx % 2 == 1:          # odd indices were inside **...**
            run.bold = True
    return para


# ---------------------------------------------------------------------------
# Line-type detection helpers
# ---------------------------------------------------------------------------

def _is_blockquote(line):
    return line.startswith('> ')


def _is_horizontal_rule(line):
    return line.strip() == '---'


def _is_article_title(line):
    return line.startswith('# ') and '【' not in line


def _is_heading2(line):
    return line.startswith('## ') or '【' in line


def _is_question(line):
    return line.startswith('??? ')


def _is_image_placeholder(line):
    return line.startswith('[img]')


# ---------------------------------------------------------------------------
# Line handlers -- each receives (doc, line, images, img_idx) and returns
# the (possibly updated) img_idx.
# ---------------------------------------------------------------------------


def _handle_article_title(doc, line):
    """Render '# text' as article title (template: 28pt Bold Red)."""
    title_text = line[2:].strip()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title_text)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.font.name = 'Arial'

def _handle_blockquote(doc, line):
    """Render '> quoted text' as italic with left indent."""
    quote_text = line[2:].strip()
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(quote_text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _handle_horizontal_rule(doc):
    """Render '---' as a thin grey separator line."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run('_' * 60)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    run.font.size = Pt(7)


def _handle_heading2(doc, line):
    """Render '## text' or 'text【section】' as Heading 1 (template style)."""
    header_text = line
    if header_text.startswith('## '):
        header_text = header_text[3:]
    heading = doc.add_heading(header_text.strip(), level=1)
    for run in heading.runs:
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def _handle_question(doc, line):
    """Render '??? interactive question' as italic blue with indent."""
    question_text = line[4:].strip()
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run('? ' + question_text)
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xcc)


def _handle_image_placeholder(doc, line, images, img_idx):
    """Render '[img]' or '[img] description' as an image or text placeholder.

    If a URL is available in *images* at *img_idx* the image is downloaded
    and embedded at 4.5 inches wide.  On any failure a centred text
    placeholder is used instead.

    Returns the incremented *img_idx*.
    """
    img_desc = line[5:].strip()
    img_url = None
    if img_idx < len(images) and images[img_idx]:
        img_url = images[img_idx]
    img_idx += 1

    inserted = False
    if img_url:
        img_data = _download_image_bytes(img_url)
        if img_data:
            try:
                suffix = '.jpg'
                if img_url.lower().endswith('.png'):
                    suffix = '.png'
                elif img_url.lower().endswith('.gif'):
                    suffix = '.gif'
                elif img_url.lower().endswith('.webp'):
                    suffix = '.webp'

                with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                    tmp.write(img_data)
                    tmp_path = tmp.name

                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(tmp_path, width=Inches(4.5))

                # Clean up the temp file as soon as the picture is embedded
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass

                inserted = True

                if img_desc:
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = caption.add_run('[Image: ' + img_desc + ']')
                    cap_run.italic = True
                    cap_run.font.size = Pt(9)
                    cap_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            except Exception:
                inserted = False

    if not inserted:
        # Text placeholder fallback
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        placeholder = '[Image'
        if img_desc:
            placeholder = placeholder + ': ' + img_desc
        placeholder = placeholder + ']'
        run = para.add_run(placeholder)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return img_idx


def _handle_regular_line(doc, line):
    """Render a normal paragraph with inline **bold** support."""
    _add_inline_formatted_paragraph(doc, line)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_docx(article, title, images, output_dir):
    """Build a styled .docx from article markdown text.

    Parameters
    ----------
    article : str
        Full article body with markdown-like formatting (see inline docs).
    title : str
        Document title, placed as centred Heading 1.
    images : list[str]
        Ordered list of image URLs referenced by ``[img]`` placeholders.
    output_dir : str
        Directory where the generated .docx will be saved.

    Returns
    -------
    str
        Bare filename (e.g. ``20260603_My_Article.docx``), not the full path.

    Raises
    ------
    Exception
        If even the minimal fallback document cannot be saved.
    """
    try:
        # ------------------------------------------------------------------
        # Primary generation path
        # ------------------------------------------------------------------
        doc = Document()

        # -- Page margins: 1 inch all sides ---------------------------------
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # -- Apply template-matched styles ----------------------------------
        _set_document_styles(doc)

        # -- Title (template: Arial 28pt Bold, dark red #C00000) -----------
        # Don't add a separate title paragraph — the article text may contain
        # a '# ' line which _handle_article_title will render correctly.
        # The `title` parameter is only used for the filename.

        # -- Parse body line by line ----------------------------------------
        img_idx = 0
        for line in article.split('\n'):
            stripped = line.strip()

            # Blank line -> empty paragraph for spacing
            if not stripped:
                doc.add_paragraph()
                continue

            if _is_article_title(stripped):
                _handle_article_title(doc, stripped)
            elif _is_blockquote(stripped):
                _handle_blockquote(doc, stripped)
            elif _is_horizontal_rule(stripped):
                _handle_horizontal_rule(doc)
            elif _is_heading2(stripped):
                _handle_heading2(doc, stripped)
            elif _is_question(stripped):
                _handle_question(doc, stripped)
            elif _is_image_placeholder(stripped):
                img_idx = _handle_image_placeholder(doc, stripped, images, img_idx)
            else:
                _handle_regular_line(doc, stripped)

        # -- Image URL appendix (always show text links at the end) --------
        doc.add_paragraph()
        sep = doc.add_paragraph()
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_run = sep.add_run('── 配图链接 ──')
        sep_run.font.size = Pt(11)
        sep_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        doc.add_paragraph()

        if images is not None and len(images) > 0:
            for i, img_url in enumerate(images):
                if not img_url:
                    continue
                # Try to embed the image inline
                img_data = _download_image_bytes(img_url)
                if img_data:
                    try:
                        suffix = '.jpg'
                        if img_url.lower().endswith('.png'):
                            suffix = '.png'
                        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                            tmp.write(img_data)
                            tmp_path = tmp.name
                        para = doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        run.add_picture(tmp_path, width=Inches(4.5))
                        try:
                            os.unlink(tmp_path)
                        except Exception:
                            pass
                    except Exception:
                        pass

                # Always show the URL as text link
                link_para = doc.add_paragraph()
                link_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                link_run = link_para.add_run('图片 ' + str(i + 1) + ': ' + img_url)
                link_run.font.size = Pt(9)
                link_run.font.color.rgb = RGBColor(0x33, 0x66, 0xcc)
                link_run.italic = True
        else:
            note = doc.add_paragraph()
            note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            n_run = note.add_run('（本文未单独配置配图）')
            n_run.font.size = Pt(9)
            n_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            n_run.italic = True

        # -- Save -----------------------------------------------------------
        filename = _build_filename(title)
        filepath = os.path.join(output_dir, filename)
        os.makedirs(output_dir, exist_ok=True)
        doc.save(filepath)

        return filename

    except Exception:
        # ------------------------------------------------------------------
        # Fallback: minimal document with title and raw text
        # ------------------------------------------------------------------
        try:
            doc = Document()
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(article)

            filename = _build_filename(title)
            filepath = os.path.join(output_dir, filename)
            os.makedirs(output_dir, exist_ok=True)
            doc.save(filepath)

            return filename
        except Exception:
            raise
